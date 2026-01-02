"""
Document Generator Service

Handles generation of PDF and DOCX documents from tailored resume content.
Supports professional formatting and multiple output formats.
"""

import re
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Literal, Optional

from logger import get_logger

# Module logger
logger = get_logger(__name__)


@dataclass
class GeneratedDocument:
    """Information about a generated document."""
    
    path: Path
    format: Literal["pdf", "docx"]
    filename: str
    size_bytes: int
    created_at: datetime
    
    def to_dict(self) -> dict:
        """Convert to dictionary."""
        return {
            "path": str(self.path),
            "format": self.format,
            "filename": self.filename,
            "size_bytes": self.size_bytes,
            "created_at": self.created_at.isoformat(),
        }


class DocumentGenerator:
    """
    Service for generating resume documents.
    
    Supports PDF and DOCX output formats with professional formatting.
    
    Usage:
        from config import settings
        generator = DocumentGenerator(settings.output_dir)
        files = generator.generate(content, ["pdf", "docx"])
    """
    
    SUPPORTED_FORMATS = {"pdf", "docx"}
    
    def __init__(self, output_dir: Path):
        """
        Initialize the document generator.
        
        Args:
            output_dir: Directory to save generated documents (from config.settings.output_dir)
        """
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        logger.info(f"Initializing DocumentGenerator with output_dir: {output_dir}")
    
    def generate(
        self,
        content: str,
        formats: list[Literal["pdf", "docx"]],
        job_title: Optional[str] = None,
        company: Optional[str] = None,
        candidate_name: Optional[str] = None,
    ) -> list[GeneratedDocument]:
        """
        Generate resume documents in specified formats.
        
        Args:
            content: Tailored resume content
            formats: List of output formats to generate
            job_title: Job title for filename
            company: Company name for filename
            candidate_name: Candidate name for document header
            
        Returns:
            list[GeneratedDocument]: Generated document information
        """
        logger.info(
            f"Generating documents in formats: {formats}",
        )
        
        # Validate formats
        invalid_formats = set(formats) - self.SUPPORTED_FORMATS
        if invalid_formats:
            logger.error(f"Invalid formats requested: {invalid_formats}")
            raise ValueError(f"Unsupported formats: {invalid_formats}")
        
        generated = []
        
        for fmt in formats:
            try:
                filename = self._generate_filename(fmt, job_title, company)
                output_path = self.output_dir / filename
                
                if fmt == "pdf":
                    self._generate_pdf(content, output_path, candidate_name)
                elif fmt == "docx":
                    self._generate_docx(content, output_path, candidate_name)
                
                # Get file size
                size_bytes = output_path.stat().st_size
                
                doc = GeneratedDocument(
                    path=output_path,
                    format=fmt,
                    filename=filename,
                    size_bytes=size_bytes,
                    created_at=datetime.now(),
                )
                generated.append(doc)
                
                logger.info(f"Generated {fmt.upper()}: {filename} ({size_bytes} bytes)")
                
            except Exception as e:
                logger.error(f"Failed to generate {fmt}: {e}")
                raise RuntimeError(f"Failed to generate {fmt.upper()} document: {e}")
        
        return generated
    
    def _generate_filename(
        self,
        format: str,
        job_title: Optional[str] = None,
        company: Optional[str] = None,
    ) -> str:
        """
        Generate a descriptive filename for the output document.
        
        Args:
            format: File format (pdf/docx)
            job_title: Optional job title
            company: Optional company name
            
        Returns:
            str: Generated filename
        """
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        parts = ["resume", "tailored"]
        
        if company:
            # Sanitize company name for filename
            safe_company = self._sanitize_for_filename(company)
            if safe_company:
                parts.append(safe_company[:30])
        
        if job_title:
            # Sanitize job title for filename
            safe_title = self._sanitize_for_filename(job_title)
            if safe_title:
                parts.append(safe_title[:30])
        
        parts.append(timestamp)
        
        filename = f"{'_'.join(parts)}.{format}"
        logger.debug(f"Generated filename: {filename}")
        
        return filename
    
    def _sanitize_for_filename(self, text: str) -> str:
        """Sanitize text for use in filename."""
        # Remove or replace invalid characters
        sanitized = re.sub(r'[<>:"/\\|?*]', '', text)
        sanitized = re.sub(r'\s+', '_', sanitized)
        sanitized = re.sub(r'_+', '_', sanitized)
        return sanitized.strip('_')
    
    def _generate_pdf(
        self, 
        content: str, 
        output_path: Path,
        candidate_name: Optional[str] = None
    ) -> None:
        """
        Generate a PDF document using FPDF2.
        
        Args:
            content: Resume content
            output_path: Output file path
            candidate_name: Optional candidate name for header
        """
        logger.debug(f"Generating PDF: {output_path}")
        
        try:
            from fpdf import FPDF
        except ImportError:
            logger.error("fpdf2 not installed. Run: pip install fpdf2")
            raise ImportError("fpdf2 is required for PDF generation. Install with: pip install fpdf2")
        
        class ResumePDF(FPDF):
            """Custom PDF class for resume formatting."""
            
            def __init__(self):
                super().__init__()
                self.set_auto_page_break(auto=True, margin=15)
            
            def header(self):
                pass  # No header needed
            
            def footer(self):
                self.set_y(-15)
                self.set_font('Helvetica', 'I', 8)
                self.set_text_color(128, 128, 128)
                self.cell(0, 10, f'Page {self.page_no()}', align='C')
        
        pdf = ResumePDF()
        pdf.add_page()
        
        # Parse and format content
        lines = content.split('\n')
        
        # Track if we're at the start (for name)
        is_first_content = True
        
        for line in lines:
            line = line.strip()
            
            if not line:
                pdf.ln(4)
                continue
            
            # Check if line is a section header (ALL CAPS or common headers)
            is_header = self._is_section_header(line)
            
            # Check for separator lines
            if self._is_separator(line):
                pdf.set_draw_color(200, 200, 200)
                pdf.line(10, pdf.get_y(), 200, pdf.get_y())
                pdf.ln(3)
                continue
            
            # Format based on line type
            if is_first_content and not is_header:
                # First line is usually the name
                pdf.set_font('Helvetica', 'B', 16)
                pdf.set_text_color(0, 0, 0)
                pdf.cell(0, 10, self._clean_text_for_pdf(line), ln=True, align='C')
                is_first_content = False
            elif is_header:
                # Section header
                pdf.ln(4)
                pdf.set_font('Helvetica', 'B', 11)
                pdf.set_text_color(44, 62, 80)
                pdf.cell(0, 8, self._clean_text_for_pdf(line), ln=True)
                pdf.set_draw_color(44, 62, 80)
                pdf.line(10, pdf.get_y(), 200, pdf.get_y())
                pdf.ln(2)
            elif line.startswith('•') or line.startswith('-') or line.startswith('*'):
                # Bullet point
                pdf.set_font('Helvetica', '', 10)
                pdf.set_text_color(0, 0, 0)
                clean_line = line.lstrip('•-* ')
                pdf.set_x(15)
                pdf.multi_cell(0, 5, f"• {self._clean_text_for_pdf(clean_line)}")
            elif '|' in line and len(line.split('|')) >= 2:
                # Contact info or inline data
                pdf.set_font('Helvetica', '', 9)
                pdf.set_text_color(80, 80, 80)
                pdf.cell(0, 6, self._clean_text_for_pdf(line), ln=True, align='C')
            else:
                # Regular text
                pdf.set_font('Helvetica', '', 10)
                pdf.set_text_color(0, 0, 0)
                pdf.multi_cell(0, 5, self._clean_text_for_pdf(line))
        
        pdf.output(output_path)
        logger.debug(f"PDF generated successfully: {output_path}")
    
    def _generate_docx(
        self, 
        content: str, 
        output_path: Path,
        candidate_name: Optional[str] = None
    ) -> None:
        """
        Generate a DOCX document using python-docx.
        
        Args:
            content: Resume content
            output_path: Output file path
            candidate_name: Optional candidate name for header
        """
        logger.debug(f"Generating DOCX: {output_path}")
        
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.style import WD_STYLE_TYPE
        except ImportError:
            logger.error("python-docx not installed. Run: pip install python-docx")
            raise ImportError("python-docx is required for DOCX generation. Install with: pip install python-docx")
        
        doc = Document()
        
        # Set up styles
        styles = doc.styles
        
        # Modify Normal style
        normal_style = styles['Normal']
        normal_style.font.name = 'Calibri'
        normal_style.font.size = Pt(11)
        
        # Parse and format content
        lines = content.split('\n')
        is_first_content = True
        
        for line in lines:
            line = line.strip()
            
            if not line:
                doc.add_paragraph()
                continue
            
            # Check if line is a section header
            is_header = self._is_section_header(line)
            
            # Skip separator lines
            if self._is_separator(line):
                continue
            
            if is_first_content and not is_header:
                # First line is usually the name
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(line)
                run.bold = True
                run.font.size = Pt(18)
                run.font.color.rgb = RGBColor(0, 0, 0)
                is_first_content = False
            elif is_header:
                # Section header
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(44, 62, 80)
                # Add bottom border effect (using a line)
                p.paragraph_format.space_after = Pt(6)
            elif line.startswith('•') or line.startswith('-') or line.startswith('*'):
                # Bullet point
                clean_line = line.lstrip('•-* ')
                p = doc.add_paragraph(clean_line, style='List Bullet')
            elif '|' in line and len(line.split('|')) >= 2:
                # Contact info
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(line)
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(80, 80, 80)
            else:
                # Regular text
                doc.add_paragraph(line)
        
        doc.save(output_path)
        logger.debug(f"DOCX generated successfully: {output_path}")
    
    def _is_section_header(self, line: str) -> bool:
        """Check if a line is a section header."""
        line_stripped = line.strip()
        
        # Common section headers
        headers = [
            'summary', 'objective', 'profile', 'experience', 'work history',
            'employment', 'education', 'skills', 'technical skills',
            'competencies', 'certifications', 'projects', 'achievements',
            'accomplishments', 'awards', 'publications', 'languages',
            'interests', 'hobbies', 'references', 'volunteer', 'leadership',
            'professional experience', 'professional summary', 'core competencies',
            'career objective', 'about me', 'contact', 'contact information'
        ]
        
        line_lower = line_stripped.lower()
        
        # Check against common headers
        for header in headers:
            if line_lower == header or line_lower.startswith(header + ':'):
                return True
        
        # Check if ALL CAPS (likely a header)
        if line_stripped.isupper() and 1 <= len(line_stripped.split()) <= 5:
            return True
        
        return False
    
    def _is_separator(self, line: str) -> bool:
        """Check if a line is a separator."""
        stripped = line.strip()
        if not stripped:
            return False
        
        # Check for lines of repeated characters
        if len(set(stripped)) == 1 and stripped[0] in '-=_*':
            return len(stripped) >= 3
        
        return False
    
    def _clean_text_for_pdf(self, text: str) -> str:
        """Clean text for PDF compatibility."""
        # Replace problematic characters
        replacements = {
            '"': '"',
            '"': '"',
            ''': "'",
            ''': "'",
            '–': '-',
            '—': '-',
            '…': '...',
            '•': '*',
        }
        
        for old, new in replacements.items():
            text = text.replace(old, new)
        
        # Remove any remaining non-latin1 characters for basic PDF compatibility
        text = text.encode('latin-1', 'replace').decode('latin-1')
        
        return text
    
    def list_generated_files(self) -> list[dict]:
        """
        List all generated files in the output directory.
        
        Returns:
            list[dict]: List of file information
        """
        files = []
        
        for path in self.output_dir.glob('resume_tailored_*'):
            if path.is_file():
                files.append({
                    "filename": path.name,
                    "format": path.suffix.lstrip('.'),
                    "size_bytes": path.stat().st_size,
                    "created_at": datetime.fromtimestamp(path.stat().st_mtime).isoformat(),
                    "path": str(path),
                })
        
        # Sort by creation time (newest first)
        files.sort(key=lambda x: x['created_at'], reverse=True)
        
        return files
    
    def cleanup_old_files(self, keep_count: int = 10) -> int:
        """
        Remove old generated files, keeping only the most recent ones.
        
        Args:
            keep_count: Number of files to keep
            
        Returns:
            int: Number of files deleted
        """
        files = list(self.output_dir.glob('resume_tailored_*'))
        
        if len(files) <= keep_count:
            return 0
        
        # Sort by modification time
        files.sort(key=lambda x: x.stat().st_mtime, reverse=True)
        
        # Delete old files
        deleted = 0
        for file in files[keep_count:]:
            try:
                file.unlink()
                deleted += 1
                logger.debug(f"Deleted old file: {file.name}")
            except Exception as e:
                logger.warning(f"Failed to delete {file.name}: {e}")
        
        if deleted > 0:
            logger.info(f"Cleaned up {deleted} old files")
        
        return deleted
