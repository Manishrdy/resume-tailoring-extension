"""
Resume Parser Service

Handles reading and parsing resume files (PDF and DOCX formats).
Extracts text content and attempts to identify resume sections.

IMPORTANT: Resume path must be passed in, obtained from config.settings.resume_path
"""

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

from logger import get_logger

# Module logger
logger = get_logger(__name__)


@dataclass
class ContactInfo:
    """Extracted contact information from resume."""
    
    email: Optional[str] = None
    phone: Optional[str] = None
    linkedin: Optional[str] = None
    github: Optional[str] = None
    location: Optional[str] = None
    name: Optional[str] = None
    
    def to_dict(self) -> dict:
        """Convert to dictionary."""
        return {
            "email": self.email,
            "phone": self.phone,
            "linkedin": self.linkedin,
            "github": self.github,
            "location": self.location,
            "name": self.name,
        }


@dataclass
class ParsedResume:
    """Structured representation of a parsed resume."""
    
    raw_text: str
    filename: str
    format: str
    sections: dict[str, str] = field(default_factory=dict)
    contact_info: Optional[ContactInfo] = None
    skills: list[str] = field(default_factory=list)
    word_count: int = 0
    
    def get_section(self, section_name: str) -> Optional[str]:
        """Get a specific section by name (case-insensitive)."""
        section_lower = section_name.lower()
        for key, value in self.sections.items():
            if section_lower in key.lower():
                return value
        return None
    
    def to_dict(self) -> dict:
        """Convert to dictionary for JSON serialization."""
        return {
            "filename": self.filename,
            "format": self.format,
            "word_count": self.word_count,
            "sections": list(self.sections.keys()),
            "skills": self.skills,
            "contact_info": self.contact_info.to_dict() if self.contact_info else None,
        }


class ResumeParser:
    """
    Service for parsing resume files.
    
    Supports PDF and DOCX formats.
    
    Usage:
        from config import settings
        parser = ResumeParser(settings.resume_path)
        resume = parser.parse()
    """
    
    SUPPORTED_FORMATS = {".pdf", ".docx", ".doc"}
    
    # Common section headers in resumes
    SECTION_PATTERNS = [
        r"(?i)^(professional\s+)?summary",
        r"(?i)^(career\s+)?objective",
        r"(?i)^profile",
        r"(?i)^about(\s+me)?",
        r"(?i)^(professional\s+)?experience",
        r"(?i)^(work\s+)?history",
        r"(?i)^employment",
        r"(?i)^education",
        r"(?i)^(technical\s+)?skills",
        r"(?i)^core\s+competencies",
        r"(?i)^competencies",
        r"(?i)^certifications?",
        r"(?i)^(notable\s+)?projects?",
        r"(?i)^achievements?",
        r"(?i)^accomplishments?",
        r"(?i)^awards?",
        r"(?i)^publications?",
        r"(?i)^languages?",
        r"(?i)^interests?",
        r"(?i)^hobbies",
        r"(?i)^references?",
        r"(?i)^volunteer(ing)?",
        r"(?i)^leadership",
        r"(?i)^training",
        r"(?i)^courses?",
    ]
    
    def __init__(self, resume_path: Path):
        """
        Initialize the parser with a resume file path.
        
        Args:
            resume_path: Path to the resume file (from config.settings.resume_path)
        """
        self.resume_path = Path(resume_path)
        logger.debug(f"Initializing ResumeParser for: {self.resume_path}")
        self._validate_file()
    
    def _validate_file(self) -> None:
        """Validate that the resume file exists and is supported."""
        logger.debug(f"Validating resume file: {self.resume_path}")
        
        if not self.resume_path.exists():
            logger.error(f"Resume file not found: {self.resume_path}")
            raise FileNotFoundError(f"Resume file not found: {self.resume_path}")
        
        if self.resume_path.suffix.lower() not in self.SUPPORTED_FORMATS:
            logger.error(f"Unsupported resume format: {self.resume_path.suffix}")
            raise ValueError(
                f"Unsupported format: {self.resume_path.suffix}. "
                f"Supported: {self.SUPPORTED_FORMATS}"
            )
        
        logger.info(f"Resume file validated: {self.resume_path.name}")
    
    def parse(self) -> ParsedResume:
        """
        Parse the resume file and extract content.
        
        Returns:
            ParsedResume: Structured resume data
        """
        logger.info(f"Parsing resume: {self.resume_path.name}")
        
        file_format = self.resume_path.suffix.lower()
        
        # Extract raw text based on format
        if file_format == ".pdf":
            raw_text = self._parse_pdf()
        elif file_format in {".docx", ".doc"}:
            raw_text = self._parse_docx()
        else:
            raise ValueError(f"Unsupported format: {file_format}")
        
        # Clean up the text
        raw_text = self._clean_text(raw_text)
        
        # Extract structured information
        sections = self._extract_sections(raw_text)
        contact_info = self._extract_contact_info(raw_text)
        skills = self._extract_skills(raw_text, sections.get("skills", ""))
        
        result = ParsedResume(
            raw_text=raw_text,
            filename=self.resume_path.name,
            format=file_format,
            sections=sections,
            contact_info=contact_info,
            skills=skills,
            word_count=len(raw_text.split()),
        )
        
        logger.info(
            f"Resume parsed successfully: {result.word_count} words, "
            f"{len(result.sections)} sections, {len(result.skills)} skills identified"
        )
        
        return result
    
    def _parse_pdf(self) -> str:
        """
        Extract text from PDF file using PyMuPDF.
        
        Returns:
            str: Extracted text content
        """
        logger.debug(f"Parsing PDF: {self.resume_path}")
        
        try:
            import fitz  # PyMuPDF
        except ImportError:
            logger.error("PyMuPDF not installed. Run: pip install pymupdf")
            raise ImportError("PyMuPDF is required for PDF parsing. Install with: pip install pymupdf")
        
        text_parts = []
        
        try:
            with fitz.open(self.resume_path) as doc:
                logger.debug(f"PDF has {len(doc)} pages")
                
                for page_num, page in enumerate(doc):
                    page_text = page.get_text("text")
                    text_parts.append(page_text)
                    logger.debug(f"Page {page_num + 1}: {len(page_text)} characters")
            
            full_text = "\n".join(text_parts)
            logger.debug(f"Total extracted: {len(full_text)} characters")
            
            return full_text
            
        except Exception as e:
            logger.error(f"Failed to parse PDF: {e}")
            raise RuntimeError(f"Failed to parse PDF file: {e}")
    
    def _parse_docx(self) -> str:
        """
        Extract text from DOCX file using python-docx.
        
        Returns:
            str: Extracted text content
        """
        logger.debug(f"Parsing DOCX: {self.resume_path}")
        
        try:
            from docx import Document
        except ImportError:
            logger.error("python-docx not installed. Run: pip install python-docx")
            raise ImportError("python-docx is required for DOCX parsing. Install with: pip install python-docx")
        
        try:
            doc = Document(self.resume_path)
            
            text_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            full_text = "\n".join(text_parts)
            logger.debug(f"Total extracted: {len(full_text)} characters")
            
            return full_text
            
        except Exception as e:
            logger.error(f"Failed to parse DOCX: {e}")
            raise RuntimeError(f"Failed to parse DOCX file: {e}")
    
    def _clean_text(self, text: str) -> str:
        """
        Clean and normalize extracted text.
        
        Args:
            text: Raw extracted text
            
        Returns:
            str: Cleaned text
        """
        # Replace multiple whitespaces with single space
        text = re.sub(r"[ \t]+", " ", text)
        
        # Replace multiple newlines with double newline
        text = re.sub(r"\n{3,}", "\n\n", text)
        
        # Strip lines
        lines = [line.strip() for line in text.split("\n")]
        text = "\n".join(lines)
        
        # Remove any null characters or other problematic chars
        text = text.replace("\x00", "")
        
        return text.strip()
    
    def _extract_sections(self, text: str) -> dict[str, str]:
        """
        Extract resume sections based on common headers.
        
        Args:
            text: Cleaned resume text
            
        Returns:
            dict: Section name -> section content
        """
        logger.debug("Extracting resume sections")
        
        sections = {}
        lines = text.split("\n")
        
        current_section = "header"
        current_content = []
        
        for line in lines:
            line_stripped = line.strip()
            
            # Check if this line is a section header
            is_header = False
            header_name = None
            
            for pattern in self.SECTION_PATTERNS:
                if re.match(pattern, line_stripped):
                    is_header = True
                    header_name = line_stripped.lower()
                    break
            
            # Also check for all-caps lines that might be headers
            if not is_header and line_stripped.isupper() and 2 <= len(line_stripped.split()) <= 5:
                is_header = True
                header_name = line_stripped.lower()
            
            if is_header and header_name:
                # Save previous section
                if current_content:
                    sections[current_section] = "\n".join(current_content).strip()
                
                # Start new section
                current_section = header_name
                current_content = []
            else:
                current_content.append(line)
        
        # Don't forget the last section
        if current_content:
            sections[current_section] = "\n".join(current_content).strip()
        
        logger.debug(f"Found {len(sections)} sections: {list(sections.keys())}")
        
        return sections
    
    def _extract_contact_info(self, text: str) -> ContactInfo:
        """
        Extract contact information from resume text.
        
        Args:
            text: Resume text
            
        Returns:
            ContactInfo: Extracted contact details
        """
        logger.debug("Extracting contact information")
        
        contact = ContactInfo()
        
        # Email pattern
        email_match = re.search(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", text)
        if email_match:
            contact.email = email_match.group()
            logger.debug(f"Found email: {contact.email}")
        
        # Phone pattern (various formats)
        phone_patterns = [
            r"\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}",  # (123) 456-7890 or 123-456-7890
            r"\+?1?[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}",  # +1 (123) 456-7890
        ]
        for pattern in phone_patterns:
            phone_match = re.search(pattern, text)
            if phone_match:
                contact.phone = phone_match.group()
                logger.debug(f"Found phone: {contact.phone}")
                break
        
        # LinkedIn pattern
        linkedin_match = re.search(r"linkedin\.com/in/[\w-]+", text, re.IGNORECASE)
        if linkedin_match:
            contact.linkedin = linkedin_match.group()
            logger.debug(f"Found LinkedIn: {contact.linkedin}")
        
        # GitHub pattern
        github_match = re.search(r"github\.com/[\w-]+", text, re.IGNORECASE)
        if github_match:
            contact.github = github_match.group()
            logger.debug(f"Found GitHub: {contact.github}")
        
        # Try to extract name (usually first non-empty line)
        lines = [l.strip() for l in text.split("\n") if l.strip()]
        if lines:
            # First line is often the name
            potential_name = lines[0]
            # Check if it looks like a name (not email, not too long, contains letters)
            if (
                "@" not in potential_name
                and len(potential_name) < 50
                and re.search(r"[a-zA-Z]", potential_name)
                and not re.match(r"^\d", potential_name)
            ):
                contact.name = potential_name
                logger.debug(f"Found name: {contact.name}")
        
        return contact
    
    def _extract_skills(self, full_text: str, skills_section: str) -> list[str]:
        """
        Extract skills from resume.
        
        Args:
            full_text: Full resume text
            skills_section: Text from skills section (if found)
            
        Returns:
            list[str]: List of identified skills
        """
        logger.debug("Extracting skills")
        
        skills = set()
        
        # Common technical skills to look for
        tech_skills = [
            # Programming languages
            r"python", r"javascript", r"typescript", r"java\b", r"c\+\+", r"c#", r"ruby", r"go\b",
            r"rust", r"php", r"swift", r"kotlin", r"scala", r"r\b", r"matlab", r"perl",
            # Web technologies
            r"html5?", r"css3?", r"react(?:\.?js)?", r"angular", r"vue(?:\.?js)?", r"node\.?js", r"express",
            r"django", r"flask", r"fastapi", r"spring\b", r"asp\.?net", r"next\.?js", r"nuxt",
            # Databases
            r"sql\b", r"mysql", r"postgresql", r"postgres", r"mongodb", r"redis", r"elasticsearch",
            r"dynamodb", r"cassandra", r"oracle", r"sqlite", r"mariadb",
            # Cloud & DevOps
            r"aws", r"amazon web services", r"azure", r"gcp", r"google cloud", 
            r"docker", r"kubernetes", r"k8s", r"terraform", r"ansible", r"jenkins", 
            r"ci/?cd", r"github actions", r"gitlab", r"circleci",
            # Data & ML
            r"machine learning", r"deep learning", r"tensorflow", r"pytorch", r"keras",
            r"pandas", r"numpy", r"scikit-learn", r"sklearn", r"nlp", r"computer vision",
            r"data science", r"data analytics", r"big data", r"spark", r"hadoop",
            # Tools & Others
            r"git\b", r"linux", r"unix", r"bash", r"shell", r"agile", r"scrum", r"jira",
            r"rest\s?api", r"graphql", r"microservices", r"api design", r"system design",
            r"oauth", r"jwt", r"websocket",
        ]
        
        # Search in skills section first, then full text
        search_text = skills_section if skills_section else full_text
        search_text_lower = search_text.lower()
        
        for skill_pattern in tech_skills:
            if re.search(rf"\b{skill_pattern}\b", search_text_lower, re.IGNORECASE):
                # Clean up the skill name for display
                skill_clean = skill_pattern.replace(r"\b", "").replace(r"\.", ".").replace(r"\s?", " ").replace(r"\+", "+").replace("?", "")
                # Normalize
                if len(skill_clean) <= 4:
                    skills.add(skill_clean.upper())
                else:
                    skills.add(skill_clean.title())
        
        skills_list = sorted(list(skills))
        logger.debug(f"Found {len(skills_list)} skills")
        
        return skills_list
