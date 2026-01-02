"""
Open-Resume Style DOCX Generator
Standalone test script - generates a professional resume in DOCX format
that is 99% identical to Open-Resume's PDF output.

Run: python test_docx_generator.py
Output: test_resume.docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


# ============================================================
# SAMPLE DATA (Matches Open-Resume demo)
# ============================================================

SAMPLE_RESUME = {
    "personal_info": {
        "name": "John Doe",
        "email": "hello@openresume.com",
        "phone": "123-456-7890",
        "location": "NYC, NY",
        "linkedin": "linkedin.com/in/johndoe",
        "github": "github.com/johndoe",
        "website": "johndoe.dev"
    },
    "professional_summary": "Software engineer with 5+ years of experience building scalable web applications and leading cross-functional teams. Passionate about creating user-centric products that drive business growth. Proven track record of delivering high-impact projects on time and mentoring junior developers.",
    "work_experiences": [
        {
            "title": "Senior Software Engineer",
            "company": "Tech Giant Inc.",
            "location": "San Francisco, CA",
            "start_date": "Jan 2021",
            "end_date": "Present",
            "bullets": [
                "Led development of microservices architecture serving 10M+ daily active users, improving system reliability by 99.9%",
                "Architected and implemented real-time data pipeline processing 1B+ events daily using Kafka and Spark",
                "Mentored team of 5 junior engineers through code reviews, pair programming, and technical workshops",
                "Reduced infrastructure costs by 40% through optimization of cloud resources and implementation of auto-scaling"
            ]
        },
        {
            "title": "Software Engineer",
            "company": "Startup Hub",
            "location": "New York, NY",
            "start_date": "Jun 2018",
            "end_date": "Dec 2020",
            "bullets": [
                "Built customer-facing dashboard using React and TypeScript, increasing user engagement by 60%",
                "Designed RESTful APIs serving 100K+ requests per minute with sub-100ms response times",
                "Implemented CI/CD pipeline reducing deployment time from 2 hours to 15 minutes",
                "Collaborated with product team to define technical requirements and deliver features iteratively"
            ]
        }
    ],
    "education": [
        {
            "degree": "Bachelor of Science in Computer Science",
            "institution": "Stanford University",
            "location": "Stanford, CA",
            "graduation_date": "May 2018",
            "gpa": "3.8/4.0",
            "highlights": []
        }
    ],
    "projects": [
        {
            "name": "Open Source Contribution - React Framework",
            "technologies": ["React", "TypeScript", "Jest"],
            "url": "github.com/johndoe/react-contrib",
            "bullets": [
                "Contributed performance optimization patches improving render time by 25%",
                "Fixed critical memory leak affecting 50K+ applications in production"
            ]
        },
        {
            "name": "AI-Powered Code Review Tool",
            "technologies": ["Python", "FastAPI", "OpenAI"],
            "url": "github.com/johndoe/ai-reviewer",
            "bullets": [
                "Built automated code review system using GPT-4 API analyzing 500+ PRs monthly",
                "Reduced code review time by 70% while maintaining code quality standards"
            ]
        }
    ],
    "skills": {
        "languages": ["Python", "JavaScript", "TypeScript", "Go", "SQL"],
        "frameworks": ["React", "Node.js", "FastAPI", "Django", "Next.js"],
        "tools": ["AWS", "Docker", "Kubernetes", "PostgreSQL", "Redis", "Kafka"],
        "other": ["System Design", "Agile/Scrum", "Technical Leadership"]
    },
    "certifications": [
        {
            "name": "AWS Solutions Architect Professional",
            "issuer": "Amazon Web Services",
            "date": "2023"
        },
        {
            "name": "Google Cloud Professional Data Engineer",
            "issuer": "Google",
            "date": "2022"
        }
    ]
}


# ============================================================
# STYLING CONFIGURATION (Matches Open-Resume)
# ============================================================

class ResumeStyle:
    """Open-Resume styling configuration."""
    
    # Page Setup
    PAGE_TOP_MARGIN = Inches(0.5)
    PAGE_BOTTOM_MARGIN = Inches(0.5)
    PAGE_LEFT_MARGIN = Inches(0.6)
    PAGE_RIGHT_MARGIN = Inches(0.6)
    
    # Fonts
    FONT_PRIMARY = "Calibri"  # ATS-safe, similar to Open-Resume's font
    
    # Font Sizes
    NAME_SIZE = Pt(20)
    CONTACT_SIZE = Pt(10)
    SECTION_HEADER_SIZE = Pt(11)
    CONTENT_SIZE = Pt(10)
    SMALL_SIZE = Pt(9)
    
    # Colors (RGB)
    COLOR_PRIMARY = RGBColor(0, 0, 0)           # Black - main text
    COLOR_SECONDARY = RGBColor(100, 100, 100)   # Gray - contact info
    COLOR_ACCENT = RGBColor(44, 62, 80)         # Dark blue - section headers
    
    # Spacing
    SECTION_SPACE_BEFORE = Pt(14)
    SECTION_SPACE_AFTER = Pt(6)
    ENTRY_SPACE_BEFORE = Pt(8)
    PARAGRAPH_SPACE_AFTER = Pt(2)
    LINE_SPACING = 1.15


# ============================================================
# DOCX GENERATOR CLASS
# ============================================================

class OpenResumeStyleGenerator:
    """Generates DOCX resume in Open-Resume style."""
    
    def __init__(self):
        self.doc = Document()
        self.style = ResumeStyle()
        self._setup_document()
    
    def _setup_document(self):
        """Configure document margins and default styles."""
        # Set margins
        for section in self.doc.sections:
            section.top_margin = self.style.PAGE_TOP_MARGIN
            section.bottom_margin = self.style.PAGE_BOTTOM_MARGIN
            section.left_margin = self.style.PAGE_LEFT_MARGIN
            section.right_margin = self.style.PAGE_RIGHT_MARGIN
        
        # Configure default paragraph style
        style = self.doc.styles['Normal']
        style.font.name = self.style.FONT_PRIMARY
        style.font.size = self.style.CONTENT_SIZE
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.line_spacing = self.style.LINE_SPACING
    
    def _add_bottom_border(self, paragraph):
        """Add a bottom border line to a paragraph (for section headers)."""
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')  # Border thickness
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '2C3E50')  # Dark blue color
        pBdr.append(bottom)
        pPr.append(pBdr)
    
    def _set_tab_stops(self, paragraph, right_tab_position=Inches(7.3)):
        """Add right-aligned tab stop for two-column layouts."""
        tab_stops = paragraph.paragraph_format.tab_stops
        tab_stops.add_tab_stop(right_tab_position, WD_TAB_ALIGNMENT.RIGHT)
    
    def add_header(self, personal_info: dict):
        """Add name and contact information header."""
        # Name - Large, Bold, Centered
        name_para = self.doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_para.paragraph_format.space_after = Pt(4)
        
        name_run = name_para.add_run(personal_info['name'])
        name_run.bold = True
        name_run.font.size = self.style.NAME_SIZE
        name_run.font.name = self.style.FONT_PRIMARY
        name_run.font.color.rgb = self.style.COLOR_PRIMARY
        
        # Contact Info - Small, Gray, Centered
        contact_parts = []
        if personal_info.get('email'):
            contact_parts.append(personal_info['email'])
        if personal_info.get('phone'):
            contact_parts.append(personal_info['phone'])
        if personal_info.get('location'):
            contact_parts.append(personal_info['location'])
        
        if contact_parts:
            contact_para = self.doc.add_paragraph()
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_para.paragraph_format.space_after = Pt(2)
            
            contact_run = contact_para.add_run(" | ".join(contact_parts))
            contact_run.font.size = self.style.CONTACT_SIZE
            contact_run.font.name = self.style.FONT_PRIMARY
            contact_run.font.color.rgb = self.style.COLOR_SECONDARY
        
        # Links - Small, Gray, Centered (second line if needed)
        link_parts = []
        if personal_info.get('linkedin'):
            link_parts.append(personal_info['linkedin'])
        if personal_info.get('github'):
            link_parts.append(personal_info['github'])
        if personal_info.get('website'):
            link_parts.append(personal_info['website'])
        
        if link_parts:
            links_para = self.doc.add_paragraph()
            links_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            links_para.paragraph_format.space_after = Pt(6)
            
            links_run = links_para.add_run(" | ".join(link_parts))
            links_run.font.size = self.style.CONTACT_SIZE
            links_run.font.name = self.style.FONT_PRIMARY
            links_run.font.color.rgb = self.style.COLOR_SECONDARY
    
    def add_section_header(self, title: str):
        """Add a section header with bottom border."""
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = self.style.SECTION_SPACE_BEFORE
        para.paragraph_format.space_after = self.style.SECTION_SPACE_AFTER
        
        run = para.add_run(title.upper())
        run.bold = True
        run.font.size = self.style.SECTION_HEADER_SIZE
        run.font.name = self.style.FONT_PRIMARY
        run.font.color.rgb = self.style.COLOR_ACCENT
        
        # Add bottom border
        self._add_bottom_border(para)
    
    def add_summary(self, summary: str):
        """Add professional summary section."""
        self.add_section_header("Professional Summary")
        
        para = self.doc.add_paragraph()
        para.paragraph_format.space_after = Pt(4)
        
        run = para.add_run(summary)
        run.font.size = self.style.CONTENT_SIZE
        run.font.name = self.style.FONT_PRIMARY
        run.font.color.rgb = self.style.COLOR_PRIMARY
    
    def add_experience(self, experiences: list):
        """Add work experience section."""
        if not experiences:
            return
        
        self.add_section_header("Experience")
        
        for i, exp in enumerate(experiences):
            # Title and Company row
            title_para = self.doc.add_paragraph()
            if i > 0:
                title_para.paragraph_format.space_before = self.style.ENTRY_SPACE_BEFORE
            else:
                title_para.paragraph_format.space_before = Pt(2)
            title_para.paragraph_format.space_after = Pt(0)
            self._set_tab_stops(title_para)
            
            # Job Title (bold)
            title_run = title_para.add_run(exp['title'])
            title_run.bold = True
            title_run.font.size = self.style.CONTENT_SIZE
            title_run.font.name = self.style.FONT_PRIMARY
            
            # Tab + Company (right-aligned)
            company_run = title_para.add_run(f"\t{exp['company']}")
            company_run.font.size = self.style.CONTENT_SIZE
            company_run.font.name = self.style.FONT_PRIMARY
            
            # Dates and Location row
            date_para = self.doc.add_paragraph()
            date_para.paragraph_format.space_after = Pt(2)
            self._set_tab_stops(date_para)
            
            # Dates (italic)
            date_run = date_para.add_run(f"{exp['start_date']} - {exp['end_date']}")
            date_run.italic = True
            date_run.font.size = self.style.SMALL_SIZE
            date_run.font.name = self.style.FONT_PRIMARY
            date_run.font.color.rgb = self.style.COLOR_SECONDARY
            
            # Tab + Location (right-aligned)
            location_run = date_para.add_run(f"\t{exp['location']}")
            location_run.font.size = self.style.SMALL_SIZE
            location_run.font.name = self.style.FONT_PRIMARY
            location_run.font.color.rgb = self.style.COLOR_SECONDARY
            
            # Bullet points
            for bullet in exp.get('bullets', []):
                self._add_bullet_point(bullet)
    
    def add_education(self, education: list):
        """Add education section."""
        if not education:
            return
        
        self.add_section_header("Education")
        
        for i, edu in enumerate(education):
            # Degree and Institution row
            degree_para = self.doc.add_paragraph()
            if i > 0:
                degree_para.paragraph_format.space_before = self.style.ENTRY_SPACE_BEFORE
            else:
                degree_para.paragraph_format.space_before = Pt(2)
            degree_para.paragraph_format.space_after = Pt(0)
            self._set_tab_stops(degree_para)
            
            # Degree (bold)
            degree_run = degree_para.add_run(edu['degree'])
            degree_run.bold = True
            degree_run.font.size = self.style.CONTENT_SIZE
            degree_run.font.name = self.style.FONT_PRIMARY
            
            # Tab + Institution (right-aligned)
            inst_run = degree_para.add_run(f"\t{edu['institution']}")
            inst_run.font.size = self.style.CONTENT_SIZE
            inst_run.font.name = self.style.FONT_PRIMARY
            
            # GPA/Date and Location row
            date_para = self.doc.add_paragraph()
            date_para.paragraph_format.space_after = Pt(2)
            self._set_tab_stops(date_para)
            
            # GPA or date info
            gpa_text = f"GPA: {edu['gpa']}" if edu.get('gpa') else edu.get('graduation_date', '')
            gpa_run = date_para.add_run(gpa_text)
            gpa_run.font.size = self.style.SMALL_SIZE
            gpa_run.font.name = self.style.FONT_PRIMARY
            gpa_run.font.color.rgb = self.style.COLOR_SECONDARY
            
            # Tab + Location or Date (right-aligned)
            loc_text = edu.get('location', '') if edu.get('gpa') else ''
            if edu.get('gpa') and edu.get('graduation_date'):
                loc_text = edu['graduation_date']
            if loc_text:
                loc_run = date_para.add_run(f"\t{loc_text}")
                loc_run.font.size = self.style.SMALL_SIZE
                loc_run.font.name = self.style.FONT_PRIMARY
                loc_run.font.color.rgb = self.style.COLOR_SECONDARY
    
    def add_projects(self, projects: list):
        """Add projects section."""
        if not projects:
            return
        
        self.add_section_header("Projects")
        
        for i, proj in enumerate(projects):
            # Project name and technologies
            proj_para = self.doc.add_paragraph()
            if i > 0:
                proj_para.paragraph_format.space_before = self.style.ENTRY_SPACE_BEFORE
            else:
                proj_para.paragraph_format.space_before = Pt(2)
            proj_para.paragraph_format.space_after = Pt(0)
            self._set_tab_stops(proj_para)
            
            # Project name (bold)
            name_run = proj_para.add_run(proj['name'])
            name_run.bold = True
            name_run.font.size = self.style.CONTENT_SIZE
            name_run.font.name = self.style.FONT_PRIMARY
            
            # Technologies (right-aligned)
            if proj.get('technologies'):
                tech_text = " | ".join(proj['technologies'])
                tech_run = proj_para.add_run(f"\t{tech_text}")
                tech_run.font.size = self.style.SMALL_SIZE
                tech_run.font.name = self.style.FONT_PRIMARY
                tech_run.font.color.rgb = self.style.COLOR_SECONDARY
            
            # URL if present
            if proj.get('url'):
                url_para = self.doc.add_paragraph()
                url_para.paragraph_format.space_after = Pt(2)
                url_run = url_para.add_run(proj['url'])
                url_run.font.size = self.style.SMALL_SIZE
                url_run.font.name = self.style.FONT_PRIMARY
                url_run.font.color.rgb = self.style.COLOR_SECONDARY
            
            # Bullet points
            for bullet in proj.get('bullets', []):
                self._add_bullet_point(bullet)
    
    def add_skills(self, skills: dict):
        """Add skills section."""
        if not skills:
            return
        
        self.add_section_header("Skills")
        
        # Combine all skills into categorized format
        skill_lines = []
        
        if skills.get('languages'):
            skill_lines.append(f"Languages: {', '.join(skills['languages'])}")
        if skills.get('frameworks'):
            skill_lines.append(f"Frameworks: {', '.join(skills['frameworks'])}")
        if skills.get('tools'):
            skill_lines.append(f"Tools: {', '.join(skills['tools'])}")
        if skills.get('other'):
            skill_lines.append(f"Other: {', '.join(skills['other'])}")
        
        for line in skill_lines:
            para = self.doc.add_paragraph()
            para.paragraph_format.space_after = Pt(2)
            
            # Split into category and skills
            if ': ' in line:
                category, skill_list = line.split(': ', 1)
                cat_run = para.add_run(f"{category}: ")
                cat_run.bold = True
                cat_run.font.size = self.style.CONTENT_SIZE
                cat_run.font.name = self.style.FONT_PRIMARY
                
                skill_run = para.add_run(skill_list)
                skill_run.font.size = self.style.CONTENT_SIZE
                skill_run.font.name = self.style.FONT_PRIMARY
    
    def add_certifications(self, certifications: list):
        """Add certifications section."""
        if not certifications:
            return
        
        self.add_section_header("Certifications")
        
        for cert in certifications:
            para = self.doc.add_paragraph()
            para.paragraph_format.space_after = Pt(2)
            self._set_tab_stops(para)
            
            # Certification name (bold)
            name_run = para.add_run(cert['name'])
            name_run.bold = True
            name_run.font.size = self.style.CONTENT_SIZE
            name_run.font.name = self.style.FONT_PRIMARY
            
            # Issuer and date (right-aligned)
            details = f"{cert.get('issuer', '')} | {cert.get('date', '')}"
            details_run = para.add_run(f"\t{details}")
            details_run.font.size = self.style.SMALL_SIZE
            details_run.font.name = self.style.FONT_PRIMARY
            details_run.font.color.rgb = self.style.COLOR_SECONDARY
    
    def _add_bullet_point(self, text: str):
        """Add a bullet point with proper formatting."""
        para = self.doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.25)
        para.paragraph_format.space_after = Pt(1)
        
        # Add bullet character
        bullet_run = para.add_run("• ")
        bullet_run.font.size = self.style.CONTENT_SIZE
        bullet_run.font.name = self.style.FONT_PRIMARY
        
        # Add text
        text_run = para.add_run(text)
        text_run.font.size = self.style.CONTENT_SIZE
        text_run.font.name = self.style.FONT_PRIMARY
    
    def generate(self, data: dict) -> Document:
        """Generate complete resume from data dictionary."""
        # Header
        self.add_header(data['personal_info'])
        
        # Summary
        if data.get('professional_summary'):
            self.add_summary(data['professional_summary'])
        
        # Experience
        if data.get('work_experiences'):
            self.add_experience(data['work_experiences'])
        
        # Education
        if data.get('education'):
            self.add_education(data['education'])
        
        # Projects
        if data.get('projects'):
            self.add_projects(data['projects'])
        
        # Skills
        if data.get('skills'):
            self.add_skills(data['skills'])
        
        # Certifications
        if data.get('certifications'):
            self.add_certifications(data['certifications'])
        
        return self.doc
    
    def save(self, filepath: str):
        """Save document to file."""
        self.doc.save(filepath)
        print(f"✅ Resume saved to: {filepath}")


# ============================================================
# MAIN - Generate Test Resume
# ============================================================

if __name__ == "__main__":
    print("=" * 50)
    print("Open-Resume Style DOCX Generator - Test")
    print("=" * 50)
    
    # Create generator
    generator = OpenResumeStyleGenerator()
    
    # Generate resume
    generator.generate(SAMPLE_RESUME)
    
    # Save to file
    output_path = "test_resume.docx"
    generator.save(output_path)
    
    print("\n📄 Open the file in Microsoft Word or Google Docs to view.")
    print("📋 The formatting should match Open-Resume's style:")
    print("   - Clean header with name and contact info")
    print("   - Section headers with blue underline")
    print("   - Two-column layout for titles/companies")
    print("   - Proper bullet points with indentation")
    print("   - Consistent spacing throughout")